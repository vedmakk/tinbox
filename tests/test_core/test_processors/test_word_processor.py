"""Tests for the Word document processor."""

from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches

from tinbox.core.processor import DocumentContent, DocumentMetadata, ProcessingError
from tinbox.core.processors.word import WordProcessor
from tinbox.core.types import FileType


@pytest.fixture
def processor() -> WordProcessor:
    """Fixture providing a Word processor instance."""
    return WordProcessor()


@pytest.fixture
def sample_docx() -> Path:
    """Fixture providing the path to the sample Word document."""
    return Path("tests/data/sample_ar.docx")


@pytest.fixture
def temp_docx(tmp_path: Path) -> Path:
    """Fixture providing a temporary Word document for testing."""
    doc = Document()
    doc.add_heading("Test Document", 0)

    # Add some regular paragraphs
    doc.add_paragraph("This is a test paragraph.")
    doc.add_paragraph("Another test paragraph.")

    # Add some Arabic text
    doc.add_paragraph("هذا نص عربي للاختبار")
    doc.add_paragraph("مزيد من النص العربي")

    # Save the document
    file_path = tmp_path / "test.docx"
    doc.save(file_path)
    return file_path


@pytest.mark.asyncio
async def test_word_processor_metadata(processor: WordProcessor, sample_docx: Path):
    """Test Word document metadata extraction."""
    metadata = await processor.get_metadata(sample_docx)

    assert isinstance(metadata, DocumentMetadata)
    assert metadata.file_type == FileType.DOCX
    assert metadata.total_pages == 1  # Single text file

    # These might be None, but should exist
    assert hasattr(metadata, "title")
    assert hasattr(metadata, "author")
    assert hasattr(metadata, "creation_date")
    assert hasattr(metadata, "modification_date")

    # Custom metadata should include Word-specific fields
    assert "category" in metadata.custom_metadata
    assert "language" in metadata.custom_metadata
    assert "contains_rtl" in metadata.custom_metadata
    assert metadata.custom_metadata["contains_rtl"]  # Should be True for Arabic text


@pytest.mark.asyncio
async def test_word_processor_content_extraction(
    processor: WordProcessor, temp_docx: Path
):
    """Test Word document content extraction."""
    content_stream = processor.extract_content(temp_docx)
    content = await content_stream.__anext__()

    assert isinstance(content, DocumentContent)
    assert content.content_type == "text/plain"
    assert content.page_number == 1  # Single text file

    # Verify content
    assert "Test Document" in content.content
    assert "This is a test paragraph" in content.content
    assert "هذا نص عربي" in content.content  # Arabic text should be present

    # Verify metadata
    assert "length" in content.metadata
    assert "contains_rtl" in content.metadata
    assert content.metadata["contains_rtl"]  # Should have RTL text


@pytest.mark.asyncio
async def test_word_processor_rtl_content(processor: WordProcessor, temp_docx: Path):
    """Test handling of right-to-left text."""
    content_stream = processor.extract_content(temp_docx)
    content = await content_stream.__anext__()

    assert "هذا نص عربي" in content.content
    assert content.metadata["contains_rtl"]


@pytest.mark.asyncio
async def test_word_processor_invalid_file(processor: WordProcessor, tmp_path: Path):
    """Test handling of invalid files."""
    # Non-existent file
    with pytest.raises(ProcessingError, match="does not exist"):
        await processor.get_metadata(tmp_path / "nonexistent.docx")

    # Invalid file type
    invalid_file = tmp_path / "test.txt"
    invalid_file.write_text("Not a Word document")
    with pytest.raises(ProcessingError, match="not supported"):
        await processor.get_metadata(invalid_file)

    # Corrupted Word document
    corrupted_file = tmp_path / "corrupted.docx"
    corrupted_file.write_text("This is not a valid DOCX file")
    with pytest.raises(ProcessingError, match="Invalid or corrupted"):
        await processor.get_metadata(corrupted_file)


# Removed test_word_processor_invalid_page_range since we no longer validate page ranges
