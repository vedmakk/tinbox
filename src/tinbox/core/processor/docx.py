"""Word document processor implementation."""

import re
from pathlib import Path
from typing import AsyncIterator, Union
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from tinbox.core.processor import (
    BaseDocumentProcessor,
    DocumentContent,
    DocumentMetadata,
    ProcessingError,
)
from tinbox.core.types import FileType
from tinbox.utils.logging import get_logger

logger = get_logger(__name__)


def detect_rtl(text: str) -> bool:
    """Detect if text contains RTL characters.

    Args:
        text: Text to check for RTL characters

    Returns:
        True if RTL characters are found, False otherwise
    """
    # RTL Unicode ranges
    rtl_ranges = [
        (0x0590, 0x05FF),  # Hebrew
        (0x0600, 0x06FF),  # Arabic
        (0x0750, 0x077F),  # Arabic Supplement
        (0x08A0, 0x08FF),  # Arabic Extended-A
        (0xFB50, 0xFDFF),  # Arabic Presentation Forms-A
        (0xFE70, 0xFEFF),  # Arabic Presentation Forms-B
    ]

    return any(
        any(start <= ord(char) <= end for start, end in rtl_ranges) for char in text
    )


def _extract_text(doc: Document) -> str:
    """Extract all text from a Word document.

    Args:
        doc: The Word document to process

    Returns:
        The document's text content as a string
    """
    paragraphs = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            paragraphs.append(paragraph.text.strip())
    return "\n".join(paragraphs)


class WordProcessor(BaseDocumentProcessor):
    """Processor for Word documents."""

    @property
    def supported_types(self) -> set[FileType]:
        """Get the file types supported by this processor."""
        return {FileType.DOCX}

    async def validate_file(self, file_path: Path) -> None:
        """Validate that the file exists and is a valid Word document.

        Args:
            file_path: Path to the Word document

        Raises:
            ProcessingError: If the file is invalid
        """
        if not file_path.exists():
            raise ProcessingError("File not found")

        if file_path.suffix.lower() != ".docx":
            raise ProcessingError("File format not supported")

    async def get_metadata(self, file_path: Path) -> DocumentMetadata:
        """Extract metadata from the Word document.

        Args:
            file_path: Path to the Word document

        Returns:
            Document metadata

        Raises:
            ProcessingError: If metadata extraction fails
        """
        await self.validate_file(file_path)

        try:
            doc = Document(file_path)
            core_props = doc.core_properties

            # Extract text and check for RTL content
            text = _extract_text(doc)

            return DocumentMetadata(
                file_type=FileType.DOCX,
                total_pages=1,  # Treat as single text file
                title=file_path.name,  # Always use filename as title
                author=core_props.author,
                creation_date=str(core_props.created) if core_props.created else None,
                modification_date=str(core_props.modified)
                if core_props.modified
                else None,
                custom_metadata={
                    "category": core_props.category,
                    "comments": core_props.comments,
                    "keywords": core_props.keywords,
                    "language": core_props.language,
                    "subject": core_props.subject,
                    "contains_rtl": detect_rtl(text),
                },
            )

        except (PackageNotFoundError, BadZipFile) as e:
            raise ProcessingError(
                f"Invalid or corrupted Word document: {str(e)}"
            ) from e
        except Exception as e:
            raise ProcessingError(f"Failed to extract Word metadata: {str(e)}") from e

    async def extract_content(
        self,
        file_path: Path,
        *,
        start_page: int = 1,
        end_page: int | None = None,
    ) -> AsyncIterator[Union[str, bytes]]:
        """Extract content from the Word document.

        Since Word documents are treated as single text files, page parameters are ignored.
        The sliding window algorithm will handle chunking the text appropriately.

        Args:
            file_path: Path to the Word document
            start_page: Must be 1 for Word files
            end_page: Must be 1 or None for Word files

        Yields:
            Text content as a single string

        Raises:
            ProcessingError: If content extraction fails
        """
        await self.validate_file(file_path)

        if start_page != 1 or (end_page is not None and end_page != 1):
            raise ProcessingError("Word files only support page 1")

        try:
            # Load document and extract text
            doc = Document(file_path)
            text = _extract_text(doc)
            yield text

        except (PackageNotFoundError, BadZipFile) as e:
            raise ProcessingError(
                f"Invalid or corrupted Word document: {str(e)}"
            ) from e
        except Exception as e:
            raise ProcessingError(f"Failed to extract Word content: {str(e)}") from e
