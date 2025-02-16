"""Word document processor implementation."""

import re
from pathlib import Path
from typing import AsyncIterator
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from tinbox.core.processor import (
    BaseDocumentProcessor,
    DocumentContent,
    DocumentMetadata,
    ProcessingError,
)
from tinbox.core.types import FileType
from tinbox.utils.logging import get_logger

logger = get_logger(__name__)


def _extract_text(doc: Document) -> str:
    """Extract all text from a Word document.
    
    Args:
        doc: The Word document to process
        
    Returns:
        The document's text content as a string
    """
    paragraphs = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            paragraphs.append(paragraph.text)
    return "\n".join(paragraphs)


class WordProcessor(BaseDocumentProcessor):
    """Processor for Word documents."""
    
    def __init__(self) -> None:
        """Initialize the Word processor."""
        super().__init__()
        self._rtl_pattern = re.compile(r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]+")
    
    @property
    def supported_types(self) -> set[FileType]:
        """Get the file types supported by this processor."""
        return {FileType.DOCX}
    
    async def get_metadata(self, file_path: Path) -> DocumentMetadata:
        """Extract metadata from the Word document.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Document metadata
            
        Raises:
            ProcessingError: If metadata extraction fails
        """
        await self.validate_file(file_path)
        
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            # Extract text and check for RTL content
            text = _extract_text(doc)
            has_rtl = bool(self._rtl_pattern.search(text))
            
            return DocumentMetadata(
                file_type=FileType.DOCX,
                total_pages=1,  # Treat as single text file
                title=core_props.title,
                author=core_props.author,
                creation_date=str(core_props.created) if core_props.created else None,
                modification_date=str(core_props.modified) if core_props.modified else None,
                custom_metadata={
                    "category": core_props.category,
                    "comments": core_props.comments,
                    "keywords": core_props.keywords,
                    "language": core_props.language,
                    "subject": core_props.subject,
                    "contains_rtl": has_rtl,
                },
            )
            
        except (PackageNotFoundError, BadZipFile) as e:
            raise ProcessingError(f"Invalid or corrupted Word document: {str(e)}") from e
        except Exception as e:
            raise ProcessingError(f"Failed to extract Word metadata: {str(e)}") from e
    
    async def extract_content(
        self,
        file_path: Path,
        *,
        start_page: int = 1,
        end_page: int | None = None,
    ) -> AsyncIterator[DocumentContent]:
        """Extract content from the Word document.
        
        Since Word documents are treated as single text files, page parameters are ignored.
        The sliding window algorithm will handle chunking the text appropriately.
        
        Args:
            file_path: Path to the Word document
            start_page: Ignored (kept for interface compatibility)
            end_page: Ignored (kept for interface compatibility)
            
        Yields:
            Document content
            
        Raises:
            ProcessingError: If content extraction fails
        """
        await self.validate_file(file_path)
        
        try:
            # Load document and extract text
            doc = Document(file_path)
            text = _extract_text(doc)
            
            # Detect if text contains RTL content
            has_rtl = bool(self._rtl_pattern.search(text))
            
            yield DocumentContent(
                content=text,
                content_type="text/plain",
                page_number=1,  # Single "page"
                metadata={
                    "length": len(text),
                    "contains_rtl": has_rtl,
                },
            )
                
        except PackageNotFoundError as e:
            raise ProcessingError(f"Invalid or corrupted Word document: {str(e)}") from e
        except Exception as e:
            raise ProcessingError(f"Failed to extract Word content: {str(e)}") from e 
